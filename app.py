import streamlit as st
import time
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from io import BytesIO
import re
import json
from pathlib import Path
from urllib.parse import urlparse

# Optional Claude support (only used if selected + key present)
try:
    import anthropic
    ANTHROPIC_SDK_AVAILABLE = True
except Exception:
    ANTHROPIC_SDK_AVAILABLE = False

# ===========================
# 🎨 PAGE CONFIGURATION
# ===========================
st.set_page_config(
    page_title="Claudio - Professional SEO Auditor",
    page_icon="👔",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ===========================
# 🔑 API CONFIGURATION
# ===========================
try:
    GEMINI_API_KEY = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=GEMINI_API_KEY)
    GEMINI_AVAILABLE = True
except Exception:
    GEMINI_AVAILABLE = False

try:
    CLAUDE_API_KEY = st.secrets.get("ANTHROPIC_API_KEY", "")
    CLAUDE_AVAILABLE = bool(CLAUDE_API_KEY) and ANTHROPIC_SDK_AVAILABLE
except Exception:
    CLAUDE_AVAILABLE = False

try:
    AHREFS_API_KEY = st.secrets.get("AHREFS_API_KEY", "")
    AHREFS_AVAILABLE = bool(AHREFS_API_KEY)
except Exception:
    AHREFS_AVAILABLE = False

# ===========================
# 🎨 CUSTOM CSS (UNCHANGED)
# ===========================
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(135deg, #2b2d42 0%, #1a1b26 100%);
    }
    
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1a1b26 0%, #121318 100%);
    }
    
    [data-testid="stMetricValue"] {
        font-size: 24px;
        color: #60a5fa;
        font-weight: 600;
    }
    
    [data-testid="stMetricLabel"] {
        color: #94a3b8;
        font-size: 13px;
        font-weight: 500;
    }
    
    .stButton>button {
        width: 100%;
        background: linear-gradient(90deg, #60a5fa 0%, #3b82f6 100%);
        color: white;
        font-weight: 600;
        border: none;
        padding: 10px 20px;
        border-radius: 6px;
        font-size: 15px;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        transform: translateY(-1px);
        box-shadow: 0 6px 12px rgba(96, 165, 250, 0.3);
    }
    
    .stTextInput>div>div>input {
        background-color: rgba(255, 255, 255, 0.05);
        color: white;
        border: 1px solid rgba(96, 165, 250, 0.3);
        border-radius: 6px;
        padding: 8px;
        font-size: 14px;
    }
    
    .stSelectbox>div>div>div {
        background-color: rgba(255, 255, 255, 0.05);
        color: white;
        border-radius: 6px;
        font-size: 14px;
    }
    
    .stRadio>div {
        background-color: rgba(255, 255, 255, 0.03);
        padding: 12px;
        border-radius: 6px;
        border: 1px solid rgba(96, 165, 250, 0.2);
    }
    
    h1 {
        color: #60a5fa;
        font-weight: 700;
    }
    
    h2, h3 {
        color: #e2e8f0;
    }
    
    .claudio-header {
        text-align: center;
        padding: 20px 0 30px 0;
        margin-bottom: 30px;
        border-bottom: 2px solid rgba(96, 165, 250, 0.2);
    }
    
    .claudio-avatar-large {
        width: 100px;
        height: 100px;
        border-radius: 50%;
        background: linear-gradient(135deg, #8B4513 0%, #654321 100%);
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 50px;
        margin: 0 auto 15px;
        border: 4px solid #60a5fa;
        box-shadow: 0 4px 12px rgba(96, 165, 250, 0.3);
    }
    
    .claudio-title {
        font-size: 42px;
        font-weight: 700;
        color: #60a5fa;
        margin: 10px 0 5px 0;
        letter-spacing: -1px;
    }
    
    .claudio-subtitle {
        font-size: 18px;
        color: #94a3b8;
        font-weight: 400;
    }
    
    .status-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 12px;
        font-weight: 600;
        margin: 4px;
    }
    
    .status-connected {
        background-color: rgba(34, 197, 94, 0.2);
        color: #22c55e;
        border: 1px solid #22c55e;
    }
    
    .status-disconnected {
        background-color: rgba(239, 68, 68, 0.2);
        color: #ef4444;
        border: 1px solid #ef4444;
    }
    
    .status-optional {
        background-color: rgba(251, 191, 36, 0.2);
        color: #fbbf24;
        border: 1px solid #fbbf24;
    }
    
    .audit-report {
        background-color: rgba(255, 255, 255, 0.03);
        padding: 30px;
        border-radius: 8px;
        border: 1px solid rgba(96, 165, 250, 0.2);
        line-height: 1.8;
    }
    
    .audit-report h1 {
        color: #60a5fa;
        border-bottom: 2px solid rgba(96, 165, 250, 0.3);
        padding-bottom: 10px;
        margin-bottom: 20px;
    }
    
    .audit-report h2 {
        color: #93c5fd;
        margin-top: 30px;
        margin-bottom: 15px;
    }
    
    .audit-report h3 {
        color: #bfdbfe;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    
    .stRadio label, .stSelectbox label {
        font-size: 13px;
        color: #94a3b8;
        font-weight: 500;
    }
</style>
""", unsafe_allow_html=True)

# ===========================
# 🧠 PATHS (templates + prompts)
# ===========================
BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
PROMPTS_DIR = BASE_DIR / "prompts"

DOCX_TEMPLATE_FULL = TEMPLATES_DIR / "SEO_Audit_Template_Full.docx"
XLSX_TEMPLATE_FULL = TEMPLATES_DIR / "SEO_Tasks_Template_Full.xlsx"
PROMPT_FULL = PROMPTS_DIR / "full.md"
PROMPT_BASIC = PROMPTS_DIR / "basic.md"

# ===========================
# 🔧 UTILS
# ===========================
def normalize_domain(url_or_domain: str) -> str:
    s = (url_or_domain or "").strip()
    if not s:
        return ""
    if s.startswith(("http://", "https://")):
        s = urlparse(s).netloc
    s = s.lower()
    if s.startswith("www."):
        s = s[4:]
    s = s.split(":")[0]
    return s

def load_prompt(path: Path) -> str:
    if path.exists():
        return path.read_text(encoding="utf-8")
    # Safe fallback so app does not crash if you forgot to commit prompts
    return ""

def strip_json_fences(text: str) -> str:
    t = (text or "").strip()
    t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*```$", "", t)
    return t.strip()

def safe_int(x, default=0):
    try:
        return int(x)
    except Exception:
        return default

def priority_from_count(cnt: int) -> str:
    cnt = safe_int(cnt, 0)
    if cnt <= 0:
        return "LOW"
    if cnt >= 50:
        return "HIGH"
    if cnt >= 10:
        return "MEDIUM"
    return "LOW"

def ahrefs_headers():
    return {
        "Authorization": f"Bearer {AHREFS_API_KEY}",
        "Accept": "application/json"
    }

def ahrefs_get(url: str, params: dict, timeout: int = 30):
    r = requests.get(url, headers=ahrefs_headers(), params=params, timeout=timeout)
    if r.status_code != 200:
        return r.status_code, None
    try:
        return r.status_code, r.json()
    except Exception:
        return r.status_code, None

# ===========================
# 🔍 WEB ANALYSIS FUNCTIONS (Basic)
# ===========================
def analyze_basic_site(url):
    """Analyzes the website extracting basic information from HTML"""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.content, 'html.parser')

        base_domain = normalize_domain(url)

        analysis = {
            'url': url,
            'status_code': response.status_code,
            'title': soup.title.string.strip() if soup.title and soup.title.string else 'No title found',
            'meta_description': '',
            'h1_tags': [],
            'h2_tags': [],
            'images_without_alt': 0,
            'total_images': 0,
            'internal_links': 0,
            'external_links': 0,
            'word_count': 0
        }

        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            analysis['meta_description'] = meta_desc.get('content', '')

        analysis['h1_tags'] = [h1.get_text(" ", strip=True) for h1 in soup.find_all('h1')]
        analysis['h2_tags'] = [h2.get_text(" ", strip=True) for h2 in soup.find_all('h2')][:5]

        images = soup.find_all('img')
        analysis['total_images'] = len(images)
        analysis['images_without_alt'] = len([img for img in images if not img.get('alt')])

        links = soup.find_all('a', href=True)
        for link in links:
            href = link['href'].strip()
            if href.startswith(('mailto:', 'tel:', '#')):
                continue
            if href.startswith('http'):
                if normalize_domain(href) != base_domain:
                    analysis['external_links'] += 1
                else:
                    analysis['internal_links'] += 1
            else:
                analysis['internal_links'] += 1

        text = soup.get_text(" ", strip=True)
        analysis['word_count'] = len(text.split())

        return analysis

    except Exception as e:
        return {'error': str(e)}

# ===========================
# 🔍 AHREFS: SITE EXPLORER (metrics/keywords/backlinks/refdomains/competitors)
# ===========================
def get_site_explorer_bundle(domain: str, country: str = "us"):
    """Minimal bundle from Site Explorer for Word placeholders + summary context."""
    bundle = {
        "metrics": {},
        "top_keywords": [],
        "refdomains": [],
        "backlinks": [],
        "competitors": []
    }

    if not AHREFS_AVAILABLE:
        return bundle

    # Metrics
    metrics_url = "https://api.ahrefs.com/v3/site-explorer/metrics"
    code, data = ahrefs_get(metrics_url, {"target": domain, "date": datetime.now().strftime("%Y-%m-%d")})
    if code == 200 and isinstance(data, dict):
        bundle["metrics"] = data.get("metrics", data)

    time.sleep(0.6)

    # Organic keywords (top 10)
    kw_url = "https://api.ahrefs.com/v3/site-explorer/organic-keywords"
    code, data = ahrefs_get(kw_url, {"target": domain, "limit": 20, "country": country, "order_by": "traffic:desc"})
    if code == 200 and isinstance(data, dict):
        kws = data.get("keywords") or data.get("data") or []
        if isinstance(kws, list):
            for it in kws[:10]:
                bundle["top_keywords"].append({
                    "keyword": it.get("keyword") or it.get("kw") or "",
                    "position": it.get("position") or it.get("pos") or "",
                    "volume": it.get("volume") or it.get("vol") or "",
                    "traffic": it.get("traffic") or it.get("traf") or "",
                    "value": it.get("traffic_value") or it.get("value") or "",
                    "url": it.get("url") or it.get("ranking_url") or ""
                })

    time.sleep(0.6)

    # Ref domains (top 2)
    ref_url = "https://api.ahrefs.com/v3/site-explorer/refdomains"
    code, data = ahrefs_get(ref_url, {"target": domain, "limit": 10, "order_by": "domain_rating:desc"})
    if code == 200 and isinstance(data, dict):
        rds = data.get("refdomains") or data.get("data") or []
        if isinstance(rds, list):
            bundle["refdomains"] = rds[:10]

    time.sleep(0.6)

    # Backlinks sample
    bl_url = "https://api.ahrefs.com/v3/site-explorer/all-backlinks"
    code, data = ahrefs_get(bl_url, {"target": domain, "limit": 10, "order_by": "domain_rating:desc"})
    if code == 200 and isinstance(data, dict):
        bls = data.get("backlinks") or data.get("data") or []
        if isinstance(bls, list):
            bundle["backlinks"] = bls[:10]

    time.sleep(0.6)

    # Competitors (top 5)
    comp_url = "https://api.ahrefs.com/v3/site-explorer/organic-competitors"
    code, data = ahrefs_get(comp_url, {"target": domain, "limit": 5, "country": country})
    if code == 200 and isinstance(data, dict):
        comps = data.get("competitors") or data.get("data") or []
        if isinstance(comps, list):
            bundle["competitors"] = comps[:5]

    return bundle

# ===========================
# 🔍 AHREFS: SITE AUDIT (projects/issues/page-explorer)
# ===========================
def site_audit_projects():
    url = "https://api.ahrefs.com/v3/site-audit/projects"
    code, data = ahrefs_get(url, {}, timeout=30)
    return code, data

def pick_project_for_domain(projects_payload: dict, domain: str):
    """
    Finds matching project for domain. Returns (project_id, project_obj) or (None, None).
    """
    if not isinstance(projects_payload, dict):
        return None, None

    projects = projects_payload.get("projects")
    if projects is None:
        # alternative shapes
        for k in ("data", "items", "result"):
            if isinstance(projects_payload.get(k), list):
                projects = projects_payload.get(k)
                break
    if not isinstance(projects, list):
        return None, None

    matches = []
    for p in projects:
        target = (p.get("target") or p.get("domain") or p.get("project_target") or "")
        if normalize_domain(target) == domain:
            ts = p.get("crawl_timestamp") or p.get("last_crawl") or p.get("updated_at") or ""
            pid = p.get("project_id") or p.get("id") or p.get("uuid")
            if pid:
                matches.append((str(ts), str(pid), p))

    if not matches:
        return None, None

    matches.sort(reverse=True, key=lambda x: x[0])
    _, pid, pobj = matches[0]
    return pid, pobj

def site_audit_issues(project_id: str):
    url = "https://api.ahrefs.com/v3/site-audit/issues"
    code, data = ahrefs_get(url, {"project_id": project_id}, timeout=30)
    return code, data

def extract_issue_list(issues_payload: dict):
    if not isinstance(issues_payload, dict):
        return []
    for k in ("issues", "data", "items", "result"):
        v = issues_payload.get(k)
        if isinstance(v, list):
            return v
    return []

def find_issue_id_and_count(issues: list, patterns: list[str]):
    best_id = None
    best_cnt = 0
    for it in issues:
        name = (it.get("name") or it.get("title") or it.get("issue_name") or "").lower()
        if not name:
            continue
        if any(p in name for p in patterns):
            iid = it.get("issue_id") or it.get("id") or it.get("uuid")
            cnt = it.get("urls_affected") or it.get("affected_urls") or it.get("affected_pages") or it.get("count") or 0
            cnt = safe_int(cnt, 0)
            if iid and cnt >= best_cnt:
                best_id = str(iid)
                best_cnt = cnt
    return best_id, best_cnt

def site_audit_page_explorer(project_id: str, issue_id: str, limit: int = 200, offset: int = 0):
    url = "https://api.ahrefs.com/v3/site-audit/page-explorer"
    params = {"project_id": project_id, "issue_id": issue_id, "limit": limit, "offset": offset}
    code, data = ahrefs_get(url, params, timeout=30)
    return code, data

def extract_page_rows(payload: dict):
    if not isinstance(payload, dict):
        return []
    for k in ("pages", "urls", "data", "items", "result"):
        v = payload.get(k)
        if isinstance(v, list):
            return v
    return []

def fetch_pages_for_issue(project_id: str, issue_id: str, max_rows: int = 300):
    rows = []
    offset = 0
    limit = min(200, max_rows)
    while len(rows) < max_rows:
        code, payload = site_audit_page_explorer(project_id, issue_id, limit=limit, offset=offset)
        if code != 200 or not isinstance(payload, dict):
            break
        chunk = extract_page_rows(payload)
        if not chunk:
            break
        rows.extend(chunk)
        if len(chunk) < limit:
            break
        offset += limit
        time.sleep(0.4)
    return rows[:max_rows]

# ===========================
# 🤖 AI (PROMPT FROM FILE + REAL MODEL ROUTING)
# ===========================
def run_llm(selected_model_label: str, prompt_text: str):
    """
    Returns dict if JSON, otherwise returns {'_raw_text': ...}
    """
    if "Gemini" in selected_model_label:
        model = genai.GenerativeModel("gemini-2.0-flash-exp")
        resp = model.generate_content(prompt_text)
        raw = strip_json_fences(getattr(resp, "text", "") or "")
    else:
        if not CLAUDE_AVAILABLE:
            return {"_raw_text": "Claude not available (missing key or SDK)."}
        client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
        # Map UI labels -> models (adjust anytime)
        claude_model = "claude-3-5-sonnet-20241022"
        if "Opus" in selected_model_label:
            claude_model = "claude-3-opus-20240229"
        msg = client.messages.create(
            model=claude_model,
            max_tokens=1400,
            temperature=0.2,
            messages=[{"role": "user", "content": prompt_text}],
        )
        raw = "".join([b.text for b in msg.content if hasattr(b, "text")])

    raw = raw.strip()
    try:
        return json.loads(strip_json_fences(raw))
    except Exception:
        return {"_raw_text": raw}

def build_prompt(audit_kind: str, context: dict):
    """
    audit_kind: 'Basic' | 'Full'
    Uses prompts/basic.md or prompts/full.md.
    Replaces {{CONTEXT_JSON}} in prompt.
    """
    prompt_path = PROMPT_BASIC if audit_kind == "Basic" else PROMPT_FULL
    p = load_prompt(prompt_path).strip()

    # If prompt file missing/empty, fallback to a safe default
    if not p:
        if audit_kind == "Basic":
            p = (
                "You are Claudio, an expert SEO auditor.\n"
                "Return ONLY JSON with keys: executive_summary, quick_wins.\n"
                "CONTEXT:\n{{CONTEXT_JSON}}\n"
            )
        else:
            p = (
                "You are Claudio, an expert SEO auditor.\n"
                "Return ONLY JSON with keys: executive_summary, content_audit_summary, technical_audit_summary, "
                "keyword_overview, backlink_observations, competitive_analysis, quick_wins.\n"
                "CONTEXT:\n{{CONTEXT_JSON}}\n"
            )

    return p.replace("{{CONTEXT_JSON}}", json.dumps(context, ensure_ascii=False, indent=2))

# ===========================
# 📄 DOCX TEMPLATE FILL (FULL)
# ===========================
PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")

def _replace_in_runs(paragraph, mapping: dict):
    # Works even if placeholders are split across runs.
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return
    changed = full_text
    for k, v in mapping.items():
        if k in changed:
            changed = changed.replace(k, str(v))
    if changed == full_text:
        return
    # write back into first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = changed
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(changed)

def _replace_in_cell(cell, mapping: dict):
    for p in cell.paragraphs:
        _replace_in_runs(p, mapping)

def _cleanup_template_instructions(doc: Document):
    """
    Removes leftover instruction text if still present after replacement.
    Conservative: only removes paragraphs containing 'e.g.' AND a placeholder marker.
    """
    for p in list(doc.paragraphs):
        t = (p.text or "")
        if "e.g." in t and "{{" in t:
            p._element.getparent().remove(p._element)

def create_word_from_full_template(site_name: str, mapping: dict) -> BytesIO:
    doc = Document(str(DOCX_TEMPLATE_FULL))

    # Replace in paragraphs
    for p in doc.paragraphs:
        _replace_in_runs(p, mapping)

    # Replace in tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_cell(cell, mapping)

    # Cleanup instructions + remove any leftover placeholders
    _cleanup_template_instructions(doc)
    for p in doc.paragraphs:
        if PLACEHOLDER_RE.search(p.text or ""):
            p.text = PLACEHOLDER_RE.sub("", p.text).strip()

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if PLACEHOLDER_RE.search(p.text or ""):
                        p.text = PLACEHOLDER_RE.sub("", p.text).strip()

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ===========================
# 📊 XLSX TEMPLATE FILL (FULL)
# ===========================
def clear_sheet_from_row(ws, start_row: int):
    maxr = ws.max_row
    if maxr >= start_row:
        ws.delete_rows(start_row, maxr - start_row + 1)

def create_excel_from_full_template(issue_rows_by_sheet: dict) -> BytesIO:
    wb = openpyxl.load_workbook(str(XLSX_TEMPLATE_FULL))
    for sheet_name, rows in issue_rows_by_sheet.items():
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        clear_sheet_from_row(ws, 2)  # keep header row
        for r in rows:
            ws.append(r)

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out

# ===========================
# 🎯 FULL AUDIT: MAP SITE AUDIT ISSUES -> YOUR XLSX SHEETS
# ===========================
ISSUE_PATTERNS = {
    "H1 Missing": ["missing h1", "h1 missing"],
    "Multiple H1": ["multiple h1", "more than one h1"],
    "Duplicate Titles": ["duplicate title"],
    "Duplicate Meta": ["duplicate meta description", "duplicate description"],
    "Title Too Long": ["title too long", "meta title too long"],
    "Title Too Short": ["title too short", "meta title too short"],
    "Meta Too Long": ["meta description too long", "description too long"],
    "Meta Too Short": ["meta description too short", "description too short"],
    "Missing Canonical": ["missing canonical"],
    "Broken Internal": ["broken internal", "internal link", "broken link"],
    "Broken External": ["broken external", "external link"],
    "Redirect Chains": ["redirect chain"],
    "Orphan Pages": ["orphan page", "orphaned page"],
    "Missing Alt Text": ["missing alt", "alt text"],
    "Broken Images": ["broken image"],
    "Thin Content": ["thin content", "low word count"],
    # Word-only / nice-to-have placeholders (counts only)
    "_ROBOTS": ["missing robots.txt", "robots.txt is missing"],
    "_SITEMAP": ["missing sitemap", "missing xml sitemap", "sitemap.xml is missing"],
    "_HTTPS": ["mixed content", "http/https", "insecure content"],
}

def row_get(row: dict, keys: list, default=""):
    for k in keys:
        if k in row and row.get(k) not in (None, ""):
            return row.get(k)
    return default

def suggest_title_from_url(url: str) -> str:
    try:
        path = urlparse(url).path.strip("/")
        if not path:
            return ""
        last = path.split("/")[-1]
        last = re.sub(r"[-_]+", " ", last).strip()
        return last.title()[:60]
    except Exception:
        return ""

def suggest_h1(title: str, url: str) -> str:
    t = (title or "").strip()
    if t and t.lower() != "missing":
        return t[:70]
    return suggest_title_from_url(url)[:70]

def suggest_meta(current_meta: str, title: str) -> str:
    m = (current_meta or "").strip()
    if m and m.lower() != "missing":
        return m[:155]
    return (title or "")[:150]

def suggest_fix_for_broken_link() -> str:
    return "Update the link to a valid destination (or remove it). If it redirects, link directly to the final URL."

def build_issue_rows_for_xlsx(project_id: str, issues_list: list, max_rows_per_sheet: int = 300):
    issue_counts = {}
    issue_rows_by_sheet = {}

    # First: resolve issue_id + counts for each sheet
    issue_ids = {}
    for sheet, pats in ISSUE_PATTERNS.items():
        iid, cnt = find_issue_id_and_count(issues_list, pats)
        issue_ids[sheet] = iid
        issue_counts[sheet] = cnt

    # Now: for each real sheet, fetch rows and format per template headers
    for sheet in [
        "H1 Missing", "Multiple H1", "Duplicate Titles", "Duplicate Meta",
        "Title Too Long", "Title Too Short", "Meta Too Long", "Meta Too Short",
        "Missing Canonical", "Broken Internal", "Broken External", "Redirect Chains",
        "Orphan Pages", "Missing Alt Text", "Broken Images", "Thin Content"
    ]:
        iid = issue_ids.get(sheet)
        cnt = issue_counts.get(sheet, 0)
        if not iid or cnt <= 0:
            issue_rows_by_sheet[sheet] = []
            continue

        rows = fetch_pages_for_issue(project_id, iid, max_rows=max_rows_per_sheet)

        formatted = []
        if sheet == "H1 Missing":
            for r in rows:
                url = row_get(r, ["url", "page_url", "address"])
                title = row_get(r, ["title", "meta_title", "page_title"])
                meta = row_get(r, ["meta_description", "description"])
                wc = row_get(r, ["word_count", "words", "content_word_count"])
                pr = "HIGH"
                formatted.append([url, title, meta, wc, pr, suggest_h1(title, url)])

        elif sheet == "Multiple H1":
            for r in rows:
                url = row_get(r, ["url", "page_url", "address"])
                h1_count = row_get(r, ["h1_count", "headings_h1_count", "count_h1"], default="")
                h1_tags = row_get(r, ["h1_tags", "h1", "headings_h1"], default="")
                pr = "HIGH"
                rec = suggest_h1(row_get(r, ["title", "meta_title", "page_title"]), url)
                formatted.append([url, h1_count, h1_tags, pr, rec])

        elif sheet == "Duplicate Titles":
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                title = row_get(r, ["title", "meta_title", "page_title"])
                dup = row_get(r, ["duplicate_count", "duplicates", "count"], default="")
                pr = "HIGH"
                sug = (title[:60] if title else suggest_title_from_url(url))
                formatted.append([url, title, dup, pr, sug])

        elif sheet == "Duplicate Meta":
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                meta = row_get(r, ["meta_description", "description"])
                dup = row_get(r, ["duplicate_count", "duplicates", "count"], default="")
                pr = "HIGH"
                sug = suggest_meta(meta, row_get(r, ["title", "meta_title", "page_title"]))
                formatted.append([url, meta, dup, pr, sug])

        elif sheet in ("Title Too Long", "Title Too Short"):
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                title = row_get(r, ["title", "meta_title", "page_title"])
                chars = row_get(r, ["character_count", "chars", "length"], default="")
                pr = "MEDIUM" if sheet == "Title Too Short" else "MEDIUM"
                sug = (title[:60] if title else suggest_title_from_url(url))
                formatted.append([url, title, chars, pr, sug])

        elif sheet in ("Meta Too Long", "Meta Too Short"):
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                meta = row_get(r, ["meta_description", "description"])
                chars = row_get(r, ["character_count", "chars", "length"], default="")
                pr = "MEDIUM"
                sug = suggest_meta(meta, row_get(r, ["title", "meta_title", "page_title"]))
                formatted.append([url, meta, chars, pr, sug])

        elif sheet == "Missing Canonical":
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                status = row_get(r, ["http_status", "status", "status_code"], default="")
                pr = "HIGH"
                canonical = url
                formatted.append([url, status, pr, canonical])

        elif sheet in ("Broken Internal", "Broken External"):
            for r in rows:
                source = row_get(r, ["source_url", "url", "page_url"])
                broken = row_get(r, ["broken_url", "link_url", "target_url"])
                status = row_get(r, ["http_status", "status", "status_code"], default="")
                anchor = row_get(r, ["anchor_text", "anchor"], default="")
                pr = "HIGH" if sheet == "Broken Internal" else "MEDIUM"
                formatted.append([source, broken, status, anchor, pr, suggest_fix_for_broken_link()])

        elif sheet == "Redirect Chains":
            for r in rows:
                initial = row_get(r, ["initial_url", "url"])
                chain = row_get(r, ["redirect_chain", "chain", "chain_path"], default="")
                final = row_get(r, ["final_url", "destination_url"], default="")
                length = row_get(r, ["chain_length", "length"], default="")
                pr = "MEDIUM"
                formatted.append([initial, chain, final, length, pr])

        elif sheet == "Orphan Pages":
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                title = row_get(r, ["title", "meta_title", "page_title"])
                wc = row_get(r, ["word_count", "words"], default="")
                incoming = row_get(r, ["incoming_links", "inlinks", "internal_inlinks"], default="")
                pr = "MEDIUM"
                action = "Add internal links from relevant hub/category pages and ensure it’s included in navigation where appropriate."
                formatted.append([url, title, wc, incoming, pr, action])

        elif sheet == "Missing Alt Text":
            for r in rows:
                page = row_get(r, ["page_url", "url"])
                img = row_get(r, ["image_url", "asset_url", "url_image"])
                pr = "LOW"
                alt = (re.sub(r"[-_]+", " ", (urlparse(img).path.split("/")[-1] if img else "")).split(".")[0]).strip().title()
                if not alt:
                    alt = suggest_title_from_url(page)[:80]
                formatted.append([page, img, pr, alt[:80]])

        elif sheet == "Broken Images":
            for r in rows:
                page = row_get(r, ["page_url", "url"])
                img = row_get(r, ["image_url", "broken_image_url", "asset_url"])
                status = row_get(r, ["http_status", "status", "status_code"], default="")
                pr = "LOW"
                fix = "Fix the image URL, restore missing asset, or remove the broken image reference."
                formatted.append([page, img, status, pr, fix])

        elif sheet == "Thin Content":
            for r in rows:
                url = row_get(r, ["url", "page_url"])
                title = row_get(r, ["title", "meta_title", "page_title"])
                wc = row_get(r, ["word_count", "words"], default="")
                pr = "MEDIUM"
                action = "Expand content to match intent; add missing sections, FAQs, examples, and improve topical depth."
                formatted.append([url, title, wc, pr, action])

        issue_rows_by_sheet[sheet] = formatted[:max_rows_per_sheet]

    return issue_counts, issue_rows_by_sheet

# ===========================
# 📄 DOCUMENT GENERATION (Basic fallback)
# ===========================
def create_word_from_content(audit_content, site_name, audit_type):
    """Creates Word document from audit content (basic mode)"""
    doc = Document()

    title = doc.add_heading(f'SEO Audit - {site_name}', 0)
    title.alignment = 1

    subtitle = doc.add_paragraph(f'{audit_type} Audit')
    subtitle.alignment = 1
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(96, 165, 250)

    doc.add_paragraph()

    lines = (audit_content or "").split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith(('- ', '* ')):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
        elif line == '---':
            doc.add_paragraph('_' * 60)
        else:
            doc.add_paragraph(line.replace('**', ''))

    doc.add_paragraph()
    footer = doc.add_paragraph(f'Generated by Claudio - {datetime.now().strftime("%B %d, %Y")}')
    footer.alignment = 1
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ===========================
# 🎨 SIDEBAR (UNCHANGED)
# ===========================
with st.sidebar:
    st.markdown("### 🏢 System Status")

    if GEMINI_AVAILABLE:
        st.markdown('<span class="status-badge status-connected">🟢 Gemini Connected</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-disconnected">🔴 Gemini Offline</span>', unsafe_allow_html=True)

    if CLAUDE_AVAILABLE:
        st.markdown('<span class="status-badge status-connected">🟢 Claude Connected</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-disconnected">🔴 Claude Offline</span>', unsafe_allow_html=True)

    if AHREFS_AVAILABLE:
        st.markdown('<span class="status-badge status-connected">🟢 Ahrefs Connected</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-optional">⚠️ Ahrefs Optional</span>', unsafe_allow_html=True)

    st.markdown("---")

    st.markdown("### ℹ️ About")
    st.markdown("""
    **Claudio** generates professional SEO audits in seconds.
    
    **Features**:
    - 🔍 Basic visual analysis
    - 💎 Full analysis with Ahrefs
    - 🤖 AI-powered insights
    - 📄 Professional reports
    """)
    st.markdown("---")
    st.caption("v2.1 - Template Edition")

# ===========================
# 🎯 MAIN INTERFACE (UNCHANGED LOOK)
# ===========================
st.markdown("""
<div class="claudio-header">
    <div class="claudio-avatar-large">👔</div>
    <div class="claudio-title">CLAUDIO</div>
    <div class="claudio-subtitle">Professional SEO Auditor</div>
</div>
""", unsafe_allow_html=True)

# Configuration
col1, col2 = st.columns([2, 1])

with col1:
    audit_type = st.radio(
        "Audit Type",
        ["🔍 Basic (Visual Analysis)", "💎 Full (With Ahrefs Data)"],
        help="Basic: Quick visual analysis\nFull: Complete with Ahrefs + Site Audit project"
    )

with col2:
    if "Full" in audit_type:
        st.info("**Full Audit**\n\n✓ Site Audit Issues (crawl)\n✓ DR / Backlinks\n✓ Keywords\n✓ Competitors")
    else:
        st.info("**Basic Audit**\n\n✓ Technical SEO\n✓ On-page analysis\n✓ Quick insights")

st.markdown("---")

# AI Model (UNCHANGED UI, but now actually used)
col1, col2 = st.columns([3, 1])

with col1:
    available_models = []
    if GEMINI_AVAILABLE:
        available_models.append("⚡ Gemini 2.0 Flash")
    if CLAUDE_AVAILABLE:
        available_models.extend(["🎯 Claude Sonnet 4.5", "👑 Claude Opus 4.5"])

    if not available_models:
        st.error("❌ No AI models configured.")
        st.stop()

    selected_model = st.selectbox("AI Model", available_models)

st.markdown("---")

# URL Input
url_input = st.text_input(
    "Website URL",
    placeholder="https://example.com",
    help="Enter the full URL including https://"
)

# Confirmation
if "Full" in audit_type:
    if AHREFS_AVAILABLE:
        st.warning("⚠️ Full Audit will use Ahrefs API credits")
        confirm_ahrefs = st.checkbox("✓ Confirm Ahrefs API usage", value=False)
    else:
        st.error("❌ Ahrefs API not configured. Cannot perform Full audit.")
        confirm_ahrefs = False
else:
    confirm_ahrefs = True

st.markdown("---")

# Generate Button
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    button_disabled = not url_input or not confirm_ahrefs

    if st.button("🚀 Generate Audit", disabled=button_disabled, use_container_width=True):
        if not url_input:
            st.error("❌ Please enter a URL")
            st.stop()

        st.markdown("---")

        progress_bar = st.progress(0)
        status_text = st.empty()

        # Step 1: Analyze site (basic always)
        status_text.text("🔍 Analyzing website...")
        progress_bar.progress(20)
        site_data = analyze_basic_site(url_input)
        time.sleep(0.6)

        if 'error' in site_data:
            st.error(f"❌ Error: {site_data['error']}")
            st.stop()

        domain = normalize_domain(url_input)
        site_name = domain or url_input.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]

        # Step 2: Full audit = Ahrefs Site Explorer + Site Audit project
        ahrefs_bundle = None
        issue_counts = {}
        issue_rows_by_sheet = {}
        competitors = []
        top_keywords = []

        if "Full" in audit_type:
            status_text.text("📊 Fetching Ahrefs data (Explorer + Site Audit)...")
            progress_bar.progress(40)

            ahrefs_bundle = get_site_explorer_bundle(domain, country="us")
            competitors = ahrefs_bundle.get("competitors", []) if isinstance(ahrefs_bundle, dict) else []
            top_keywords = ahrefs_bundle.get("top_keywords", []) if isinstance(ahrefs_bundle, dict) else []

            # Resolve Site Audit project
            code, projects_payload = site_audit_projects()
            if code != 200 or not projects_payload:
                st.error("❌ Could not fetch Ahrefs Site Audit projects. Check API key/plan.")
                st.stop()

            project_id, _proj = pick_project_for_domain(projects_payload, domain)
            if not project_id:
                st.error(f"❌ No Site Audit project found for {domain}. Make sure it exists in Ahrefs Site Audit.")
                st.stop()

            # Fetch issues
            code, issues_payload = site_audit_issues(project_id)
            if code != 200 or not issues_payload:
                st.error("❌ Could not fetch Site Audit issues for this project.")
                st.stop()

            issues_list = extract_issue_list(issues_payload)

            # Build rows for Excel + counts for Word placeholders
            issue_counts, issue_rows_by_sheet = build_issue_rows_for_xlsx(
                project_id=project_id,
                issues_list=issues_list,
                max_rows_per_sheet=300
            )

            time.sleep(0.6)

        # Step 3: Generate AI narrative blocks from external prompt
        status_text.text("🤖 Generating audit content...")
        progress_bar.progress(60)

        type_audit = "Basic" if "Basic" in audit_type else "Full"

        context = {
            "domain": domain,
            "audit_date": datetime.now().strftime("%B %Y"),
            "basic_onpage": site_data,
            "ahrefs": ahrefs_bundle or {},
            "site_audit_issue_counts": issue_counts or {},
            "competitors": competitors or [],
            "top_keywords": top_keywords or [],
        }

        prompt_text = build_prompt(type_audit, context)
        ai_out = run_llm(selected_model, prompt_text)

        # Provide a preview markdown (keep the same preview tab behavior)
        if isinstance(ai_out, dict) and "_raw_text" not in ai_out:
            # Build readable markdown from JSON keys if present
            audit_content = ""
            if type_audit == "Full":
                audit_content += "## EXECUTIVE SUMMARY\n"
                audit_content += (ai_out.get("executive_summary", "") or "") + "\n\n"
                audit_content += "## CONTENT AUDIT\n"
                audit_content += (ai_out.get("content_audit_summary", "") or "") + "\n\n"
                audit_content += "## TECHNICAL AUDIT\n"
                audit_content += (ai_out.get("technical_audit_summary", "") or "") + "\n\n"
                audit_content += "## KEYWORD PERFORMANCE\n"
                audit_content += (ai_out.get("keyword_overview", "") or "") + "\n\n"
                audit_content += "## BACKLINK PROFILE\n"
                audit_content += (ai_out.get("backlink_observations", "") or "") + "\n\n"
                audit_content += "## COMPETITIVE ANALYSIS\n"
                audit_content += (ai_out.get("competitive_analysis", "") or "") + "\n\n"
                audit_content += "## QUICK WINS\n"
                qws = ai_out.get("quick_wins", []) or []
                for i, qw in enumerate(qws[:8], start=1):
                    action = (qw.get("action") if isinstance(qw, dict) else str(qw)) or ""
                    impact = qw.get("impact", "") if isinstance(qw, dict) else ""
                    effort = qw.get("effort", "") if isinstance(qw, dict) else ""
                    audit_content += f"{i}. {action} (Impact: {impact}, Effort: {effort})\n"
            else:
                audit_content += "## EXECUTIVE SUMMARY\n"
                audit_content += (ai_out.get("executive_summary", "") or "") + "\n\n"
                audit_content += "## QUICK WINS\n"
                qws = ai_out.get("quick_wins", []) or []
                for i, qw in enumerate(qws[:8], start=1):
                    action = (qw.get("action") if isinstance(qw, dict) else str(qw)) or ""
                    impact = qw.get("impact", "") if isinstance(qw, dict) else ""
                    effort = qw.get("effort", "") if isinstance(qw, dict) else ""
                    audit_content += f"{i}. {action} (Impact: {impact}, Effort: {effort})\n"
        else:
            audit_content = ai_out.get("_raw_text", "AI output not available.")

        # Step 4: Create documents (FULL uses templates; BASIC uses markdown-to-docx)
        status_text.text("📄 Creating documents...")
        progress_bar.progress(80)

        doc_file = None
        excel_file = None

        if type_audit == "Full":
            # Build placeholder mapping for the Word template you uploaded
            metrics = (ahrefs_bundle or {}).get("metrics", {}) if isinstance(ahrefs_bundle, dict) else {}
            dr = metrics.get("domain_rating", metrics.get("dr", 0))
            ar = metrics.get("ahrefs_rank", metrics.get("rank", ""))
            backlinks_total = metrics.get("backlinks", 0)
            dofollow_backlinks = metrics.get("dofollow_backlinks", metrics.get("dofollow", ""))
            refdomains_total = metrics.get("refdomains", metrics.get("referring_domains", 0))
            dofollow_refdomains = metrics.get("dofollow_refdomains", "")

            organic_keywords = metrics.get("organic_keywords", 0)
            organic_traffic = metrics.get("organic_traffic", 0)

            # keyword distribution (from top 100 if present)
            dist = {"1_3": 0, "4_10": 0, "11_20": 0, "21_50": 0, "51_100": 0}
            # If ahrefs returned more keywords in bundle, we only have top 10. Keep distribution zeros (safe).
            # You can extend later if you want the full distribution.

            # Take KW_1 and KW_2 for the template
            kw1 = top_keywords[0] if len(top_keywords) > 0 else {}
            kw2 = top_keywords[1] if len(top_keywords) > 1 else {}

            # Take REF_1 and REF_2 for the template
            refdomains = (ahrefs_bundle or {}).get("refdomains", []) if isinstance(ahrefs_bundle, dict) else []
            ref1 = refdomains[0] if len(refdomains) > 0 else {}
            ref2 = refdomains[1] if len(refdomains) > 1 else {}

            # Competitors (COMP_1..5)
            comps = competitors or []

            # AI narrative placeholders
            exec_sum = ai_out.get("executive_summary", "") if isinstance(ai_out, dict) else ""
            content_sum = ai_out.get("content_audit_summary", "") if isinstance(ai_out, dict) else ""
            technical_sum = ai_out.get("technical_audit_summary", "") if isinstance(ai_out, dict) else ""
            keyword_overview = ai_out.get("keyword_overview", "") if isinstance(ai_out, dict) else ""
            backlink_obs = ai_out.get("backlink_observations", "") if isinstance(ai_out, dict) else ""
            competitive_analysis = ai_out.get("competitive_analysis", "") if isinstance(ai_out, dict) else ""
            quick_wins = ai_out.get("quick_wins", []) if isinstance(ai_out, dict) else []

            # Counts from Site Audit issues for Word placeholders (match your template names)
            mapping = {
                "{{DOMAIN}}": site_name,
                "{{AUDIT_DATE}}": datetime.now().strftime("%B %Y"),

                "{{DOMAIN_RATING}}": dr,
                "{{AHREFS_RANK}}": ar,
                "{{REFERRING_DOMAINS}}": refdomains_total,
                "{{ORGANIC_KEYWORDS}}": organic_keywords,
                "{{ORGANIC_TRAFFIC}}": organic_traffic,
                "{{TOTAL_BACKLINKS}}": backlinks_total,
                "{{DOFOLLOW_BACKLINKS}}": dofollow_backlinks,
                "{{DOFOLLOW_REFDOMAINS}}": dofollow_refdomains,

                "{{EXECUTIVE_SUMMARY}}": exec_sum,
                "{{CONTENT_AUDIT_SUMMARY}}": content_sum,
                "{{TECHNICAL_AUDIT_SUMMARY}}": technical_sum,
                "{{KEYWORD_OVERVIEW}}": keyword_overview,
                "{{BACKLINK_OBSERVATIONS}}": backlink_obs,
                "{{COMPETITIVE_ANALYSIS}}": competitive_analysis,

                "{{TOP_KEYWORDS_COUNT}}": len(top_keywords),

                "{{KW_POS_1_3}}": dist["1_3"],
                "{{KW_POS_4_10}}": dist["4_10"],
                "{{KW_POS_11_20}}": dist["11_20"],
                "{{KW_POS_21_50}}": dist["21_50"],
                "{{KW_POS_51_100}}": dist["51_100"],

                "{{KW_1}}": kw1.get("keyword", ""),
                "{{KW_1_POS}}": kw1.get("position", ""),
                "{{KW_1_VOL}}": kw1.get("volume", ""),
                "{{KW_1_TRAFFIC}}": kw1.get("traffic", ""),
                "{{KW_1_VALUE}}": kw1.get("value", ""),
                "{{KW_1_URL}}": kw1.get("url", ""),

                "{{KW_2}}": kw2.get("keyword", ""),
                "{{KW_2_POS}}": kw2.get("position", ""),
                "{{KW_2_VOL}}": kw2.get("volume", ""),
                "{{KW_2_TRAFFIC}}": kw2.get("traffic", ""),
                "{{KW_2_VALUE}}": kw2.get("value", ""),
                "{{KW_2_URL}}": kw2.get("url", ""),

                "{{REF_1_DOMAIN}}": ref1.get("domain", ref1.get("refdomain", "")),
                "{{REF_1_DR}}": ref1.get("domain_rating", ref1.get("dr", "")),
                "{{REF_1_LINKS}}": ref1.get("links", ref1.get("backlinks", "")),
                "{{REF_1_DF}}": ref1.get("dofollow_links", ref1.get("dofollow", "")),
                "{{REF_1_TRAFFIC}}": ref1.get("traffic", ref1.get("organic_traffic", "")),

                "{{REF_2_DOMAIN}}": ref2.get("domain", ref2.get("refdomain", "")),
                "{{REF_2_DR}}": ref2.get("domain_rating", ref2.get("dr", "")),
                "{{REF_2_LINKS}}": ref2.get("links", ref2.get("backlinks", "")),
                "{{REF_2_DF}}": ref2.get("dofollow_links", ref2.get("dofollow", "")),
                "{{REF_2_TRAFFIC}}": ref2.get("traffic", ref2.get("organic_traffic", "")),

                "{{MISSING_H1_COUNT}}": issue_counts.get("H1 Missing", 0),
                "{{MULTIPLE_H1_COUNT}}": issue_counts.get("Multiple H1", 0),
                "{{DUP_TITLES_COUNT}}": issue_counts.get("Duplicate Titles", 0),
                "{{DUP_META_COUNT}}": issue_counts.get("Duplicate Meta", 0),
                "{{TITLE_LONG_COUNT}}": issue_counts.get("Title Too Long", 0),
                "{{TITLE_SHORT_COUNT}}": issue_counts.get("Title Too Short", 0),
                "{{META_LONG_COUNT}}": issue_counts.get("Meta Too Long", 0),
                "{{META_SHORT_COUNT}}": issue_counts.get("Meta Too Short", 0),
                "{{THIN_CONTENT_COUNT}}": issue_counts.get("Thin Content", 0),
                "{{MISSING_ALT_COUNT}}": issue_counts.get("Missing Alt Text", 0),
                "{{BROKEN_IMAGES_COUNT}}": issue_counts.get("Broken Images", 0),

                "{{MISSING_CANONICAL_COUNT}}": issue_counts.get("Missing Canonical", 0),
                "{{BROKEN_INTERNAL_COUNT}}": issue_counts.get("Broken Internal", 0),
                "{{BROKEN_EXTERNAL_COUNT}}": issue_counts.get("Broken External", 0),
                "{{REDIRECT_CHAINS_COUNT}}": issue_counts.get("Redirect Chains", 0),
                "{{ORPHAN_PAGES_COUNT}}": issue_counts.get("Orphan Pages", 0),

                "{{MISSING_CANONICAL_PRIORITY}}": priority_from_count(issue_counts.get("Missing Canonical", 0)),
                "{{BROKEN_INTERNAL_PRIORITY}}": priority_from_count(issue_counts.get("Broken Internal", 0)),
                "{{BROKEN_EXTERNAL_PRIORITY}}": priority_from_count(issue_counts.get("Broken External", 0)),
                "{{REDIRECT_CHAINS_PRIORITY}}": priority_from_count(issue_counts.get("Redirect Chains", 0)),
                "{{ORPHAN_PAGES_PRIORITY}}": priority_from_count(issue_counts.get("Orphan Pages", 0)),

                "{{TITLE_LONG_PRIORITY}}": priority_from_count(issue_counts.get("Title Too Long", 0)),
                "{{TITLE_SHORT_PRIORITY}}": priority_from_count(issue_counts.get("Title Too Short", 0)),
                "{{META_LONG_PRIORITY}}": priority_from_count(issue_counts.get("Meta Too Long", 0)),
                "{{META_SHORT_PRIORITY}}": priority_from_count(issue_counts.get("Meta Too Short", 0)),
                "{{THIN_CONTENT_PRIORITY}}": priority_from_count(issue_counts.get("Thin Content", 0)),
                "{{MISSING_ALT_PRIORITY}}": priority_from_count(issue_counts.get("Missing Alt Text", 0)),
                "{{BROKEN_IMAGES_PRIORITY}}": priority_from_count(issue_counts.get("Broken Images", 0)),

                # overview counts
                "{{CONTENT_ISSUES_COUNT}}": sum([
                    issue_counts.get("H1 Missing", 0),
                    issue_counts.get("Multiple H1", 0),
                    issue_counts.get("Duplicate Titles", 0),
                    issue_counts.get("Duplicate Meta", 0),
                    issue_counts.get("Title Too Long", 0),
                    issue_counts.get("Title Too Short", 0),
                    issue_counts.get("Meta Too Long", 0),
                    issue_counts.get("Meta Too Short", 0),
                    issue_counts.get("Thin Content", 0),
                    issue_counts.get("Missing Alt Text", 0),
                    issue_counts.get("Broken Images", 0),
                ]),
                "{{TECHNICAL_ISSUES_COUNT}}": sum([
                    issue_counts.get("Missing Canonical", 0),
                    issue_counts.get("Broken Internal", 0),
                    issue_counts.get("Broken External", 0),
                    issue_counts.get("Redirect Chains", 0),
                    issue_counts.get("Orphan Pages", 0),
                    issue_counts.get("_ROBOTS", 0),
                    issue_counts.get("_SITEMAP", 0),
                    issue_counts.get("_HTTPS", 0),
                ]),
                "{{CONTENT_PRIORITY}}": "HIGH" if sum(issue_counts.get(k, 0) for k in [
                    "H1 Missing","Duplicate Titles","Duplicate Meta"
                ]) > 0 else "MEDIUM",
                "{{TECHNICAL_PRIORITY}}": "HIGH" if sum(issue_counts.get(k, 0) for k in [
                    "Missing Canonical","Broken Internal"
                ]) > 0 else "MEDIUM",

                "{{BACKLINK_OPP_COUNT}}": refdomains_total,
                "{{COMPETITIVE_GAPS_COUNT}}": len(comps),

                # your row in competitor table
                "{{YOUR_DR}}": dr,
                "{{YOUR_REFDOM}}": refdomains_total,
                "{{YOUR_KW}}": organic_keywords,
                "{{YOUR_TRAFFIC}}": organic_traffic,
                "{{YOUR_VALUE}}": "",
            }

            # quick wins placeholders (your template expects QUICK_WIN_1..5 and QW1_IMPACT/EFFORT etc.)
            if not isinstance(quick_wins, list):
                quick_wins = []
            while len(quick_wins) < 5:
                quick_wins.append({"action": "", "impact": "Medium", "effort": "Low"})

            for i in range(5):
                qw = quick_wins[i] if isinstance(quick_wins[i], dict) else {"action": str(quick_wins[i])}
                mapping[f"{{{{QUICK_WIN_{i+1}}}}}"] = qw.get("action", "")
                mapping[f"{{{{QW{i+1}_IMPACT}}}}"] = qw.get("impact", "Medium")
                mapping[f"{{{{QW{i+1}_EFFORT}}}}"] = qw.get("effort", "Low")

            # competitor placeholders COMP_1..5 (domain + DR/RD/KW/Traffic/Value)
            for i in range(5):
                c = comps[i] if i < len(comps) else {}
                mapping[f"{{{{COMP_{i+1}}}}}"] = c.get("domain", c.get("target", ""))
                mapping[f"{{{{COMP_{i+1}_DR}}}}"] = c.get("domain_rating", c.get("dr", ""))
                mapping[f"{{{{COMP_{i+1}_REFDOM}}}}"] = c.get("refdomains", c.get("referring_domains", ""))
                mapping[f"{{{{COMP_{i+1}_KW}}}}"] = c.get("organic_keywords", c.get("keywords", ""))
                mapping[f"{{{{COMP_{i+1}_TRAFFIC}}}}"] = c.get("organic_traffic", c.get("traffic", ""))
                mapping[f"{{{{COMP_{i+1}_VALUE}}}}"] = c.get("traffic_value", c.get("value", ""))

            # Generate Word from your template
            if not DOCX_TEMPLATE_FULL.exists():
                st.error("❌ Missing Word template: templates/SEO_Audit_Template_Full.docx")
                st.stop()
            doc_file = create_word_from_full_template(site_name, mapping)

            # Generate Excel from your template (fill only issue sheets; keep SEO Tasks as-is)
            if not XLSX_TEMPLATE_FULL.exists():
                st.error("❌ Missing Excel template: templates/SEO_Tasks_Template_Full.xlsx")
                st.stop()
            excel_file = create_excel_from_full_template(issue_rows_by_sheet)

        else:
            # Basic: keep old behavior (simple doc from markdown content)
            doc_file = create_word_from_content(audit_content, site_name, type_audit)
            excel_file = None

        progress_bar.progress(100)
        status_text.text("✅ Complete!")
        time.sleep(0.5)

        progress_bar.empty()
        status_text.empty()

        # Results
        st.markdown("---")
        st.success("✅ Audit completed successfully!")

        tab1, tab2 = st.tabs(["📄 Preview", "📥 Download"])

        with tab1:
            st.markdown('<div class="audit-report">', unsafe_allow_html=True)
            st.markdown(audit_content)
            st.markdown('</div>', unsafe_allow_html=True)

        with tab2:
            st.markdown("### Download Your Documents")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### 📄 Audit Report")
                st.download_button(
                    label="📥 Download Report (.docx)",
                    data=doc_file,
                    file_name=f"SEO_Audit_{site_name}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            with col2:
                if excel_file and type_audit == "Full":
                    st.markdown("#### 📊 Task List")
                    st.download_button(
                        label="📥 Download Tasks (.xlsx)",
                        data=excel_file,
                        file_name=f"SEO_Tasks_{site_name}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                else:
                    st.info("📊 Task list only for Full audits")

# Footer (UNCHANGED)
st.markdown("---")
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("**Claudio SEO Auditor**")
    st.caption("Professional audits in seconds")

with col2:
    st.markdown("**Powered by**")
    st.caption("Anthropic • Google • Ahrefs")

with col3:
    st.markdown("**Need help?**")
    st.caption("[Documentation](#) • [Support](#)")
