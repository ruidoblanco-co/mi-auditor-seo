import streamlit as st
import time
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import google.generativeai as genai
from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from io import BytesIO
import re

# ===========================
# 🎨 PAGE CONFIGURATION
# ===========================
st.set_page_config(
    page_title="Claudio - Professional SEO Auditor",
    page_icon="👔",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ===========================
# 🔑 API CONFIGURATION
# ===========================
try:
    GEMINI_API_KEY = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=GEMINI_API_KEY)
    GEMINI_AVAILABLE = True
except Exception as e:
    GEMINI_AVAILABLE = False

try:
    CLAUDE_API_KEY = st.secrets.get("ANTHROPIC_API_KEY", "")
    CLAUDE_AVAILABLE = bool(CLAUDE_API_KEY)
except:
    CLAUDE_AVAILABLE = False

try:
    AHREFS_API_KEY = st.secrets.get("AHREFS_API_KEY", "")
    AHREFS_AVAILABLE = bool(AHREFS_API_KEY)
except:
    AHREFS_AVAILABLE = False

# ===========================
# 🎨 CUSTOM CSS
# ===========================
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(135deg, #2b2d42 0%, #1a1b26 100%);
    }
    
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1a1b26 0%, #121318 100%);
    }
    
    [data-testid="stMetricValue"] {
        font-size: 24px;
        color: #60a5fa;
        font-weight: 600;
    }
    
    [data-testid="stMetricLabel"] {
        color: #94a3b8;
        font-size: 13px;
        font-weight: 500;
    }
    
    .stButton>button {
        width: 100%;
        background: linear-gradient(90deg, #60a5fa 0%, #3b82f6 100%);
        color: white;
        font-weight: 600;
        border: none;
        padding: 10px 20px;
        border-radius: 6px;
        font-size: 15px;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        transform: translateY(-1px);
        box-shadow: 0 6px 12px rgba(96, 165, 250, 0.3);
    }
    
    .stTextInput>div>div>input {
        background-color: rgba(255, 255, 255, 0.05);
        color: white;
        border: 1px solid rgba(96, 165, 250, 0.3);
        border-radius: 6px;
        padding: 8px;
        font-size: 14px;
    }
    
    .stSelectbox>div>div>div {
        background-color: rgba(255, 255, 255, 0.05);
        color: white;
        border-radius: 6px;
        font-size: 14px;
    }
    
    .stRadio>div {
        background-color: rgba(255, 255, 255, 0.03);
        padding: 12px;
        border-radius: 6px;
        border: 1px solid rgba(96, 165, 250, 0.2);
    }
    
    h1 {
        color: #60a5fa;
        font-weight: 700;
    }
    
    h2, h3 {
        color: #e2e8f0;
    }
    
    .claudio-header {
        text-align: center;
        padding: 20px 0 30px 0;
        margin-bottom: 30px;
        border-bottom: 2px solid rgba(96, 165, 250, 0.2);
    }
    
    .claudio-avatar-large {
        width: 100px;
        height: 100px;
        border-radius: 50%;
        background: linear-gradient(135deg, #8B4513 0%, #654321 100%);
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 50px;
        margin: 0 auto 15px;
        border: 4px solid #60a5fa;
        box-shadow: 0 4px 12px rgba(96, 165, 250, 0.3);
    }
    
    .claudio-title {
        font-size: 42px;
        font-weight: 700;
        color: #60a5fa;
        margin: 10px 0 5px 0;
        letter-spacing: -1px;
    }
    
    .claudio-subtitle {
        font-size: 18px;
        color: #94a3b8;
        font-weight: 400;
    }
    
    .status-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 12px;
        font-weight: 600;
        margin: 4px;
    }
    
    .status-connected {
        background-color: rgba(34, 197, 94, 0.2);
        color: #22c55e;
        border: 1px solid #22c55e;
    }
    
    .status-disconnected {
        background-color: rgba(239, 68, 68, 0.2);
        color: #ef4444;
        border: 1px solid #ef4444;
    }
    
    .status-optional {
        background-color: rgba(251, 191, 36, 0.2);
        color: #fbbf24;
        border: 1px solid #fbbf24;
    }
    
    .audit-report {
        background-color: rgba(255, 255, 255, 0.03);
        padding: 30px;
        border-radius: 8px;
        border: 1px solid rgba(96, 165, 250, 0.2);
        line-height: 1.8;
    }
    
    .audit-report h1 {
        color: #60a5fa;
        border-bottom: 2px solid rgba(96, 165, 250, 0.3);
        padding-bottom: 10px;
        margin-bottom: 20px;
    }
    
    .audit-report h2 {
        color: #93c5fd;
        margin-top: 30px;
        margin-bottom: 15px;
    }
    
    .audit-report h3 {
        color: #bfdbfe;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    
    .stRadio label, .stSelectbox label {
        font-size: 13px;
        color: #94a3b8;
        font-weight: 500;
    }
</style>
""", unsafe_allow_html=True)

# ===========================
# 🔍 AHREFS API FUNCTIONS
# ===========================
def get_ahrefs_data(domain):
    """Get comprehensive data from Ahrefs API"""
    
    if not AHREFS_AVAILABLE:
        return None
    
    headers = {
        "Authorization": f"Bearer {AHREFS_API_KEY}",
        "Accept": "application/json"
    }
    
    # Clean domain
    domain = domain.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
    
    ahrefs_data = {
        'domain': domain,
        'domain_rating': 0,
        'url_rating': 0,
        'backlinks': 0,
        'referring_domains': 0,
        'organic_keywords': 0,
        'organic_traffic': 0,
        'top_keywords': [],
        'top_pages': [],
        'backlink_sample': []
    }
    
    try:
        # Domain metrics
        metrics_url = "https://api.ahrefs.com/v3/site-explorer/metrics"
        metrics_params = {
            "target": domain,
            "date": datetime.now().strftime("%Y-%m-%d")
        }
        
        response = requests.get(metrics_url, headers=headers, params=metrics_params, timeout=30)
        if response.status_code == 200:
            data = response.json()
            if 'metrics' in data:
                m = data['metrics']
                ahrefs_data['domain_rating'] = m.get('domain_rating', 0)
                ahrefs_data['url_rating'] = m.get('url_rating', 0)
                ahrefs_data['backlinks'] = m.get('backlinks', 0)
                ahrefs_data['referring_domains'] = m.get('refdomains', 0)
                ahrefs_data['organic_keywords'] = m.get('organic_keywords', 0)
                ahrefs_data['organic_traffic'] = m.get('organic_traffic', 0)
        
        time.sleep(1)
        
        # Top keywords
        keywords_url = "https://api.ahrefs.com/v3/site-explorer/organic-keywords"
        keywords_params = {
            "target": domain,
            "limit": 20,
            "order_by": "volume:desc"
        }
        
        response = requests.get(keywords_url, headers=headers, params=keywords_params, timeout=30)
        if response.status_code == 200:
            data = response.json()
            if 'keywords' in data:
                ahrefs_data['top_keywords'] = data['keywords'][:10]
        
        time.sleep(1)
        
        # Top pages
        pages_url = "https://api.ahrefs.com/v3/site-explorer/top-pages"
        pages_params = {
            "target": domain,
            "limit": 10
        }
        
        response = requests.get(pages_url, headers=headers, params=pages_params, timeout=30)
        if response.status_code == 200:
            data = response.json()
            if 'pages' in data:
                ahrefs_data['top_pages'] = data['pages'][:5]
        
        time.sleep(1)
        
        # Backlink sample
        backlinks_url = "https://api.ahrefs.com/v3/site-explorer/all-backlinks"
        backlinks_params = {
            "target": domain,
            "limit": 10,
            "order_by": "domain_rating:desc"
        }
        
        response = requests.get(backlinks_url, headers=headers, params=backlinks_params, timeout=30)
        if response.status_code == 200:
            data = response.json()
            if 'backlinks' in data:
                ahrefs_data['backlink_sample'] = data['backlinks'][:5]
        
        return ahrefs_data
        
    except Exception as e:
        st.warning(f"⚠️ Ahrefs API error: {str(e)}")
        return ahrefs_data  # Return with defaults

# ===========================
# 🔍 WEB ANALYSIS FUNCTIONS
# ===========================
def analyze_basic_site(url):
    """Analyzes the website extracting basic information from HTML"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        analysis = {
            'url': url,
            'status_code': response.status_code,
            'title': soup.title.string if soup.title else 'No title found',
            'meta_description': '',
            'h1_tags': [],
            'h2_tags': [],
            'images_without_alt': 0,
            'total_images': 0,
            'internal_links': 0,
            'external_links': 0,
            'word_count': 0
        }
        
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc:
            analysis['meta_description'] = meta_desc.get('content', '')
        
        analysis['h1_tags'] = [h1.get_text().strip() for h1 in soup.find_all('h1')]
        analysis['h2_tags'] = [h2.get_text().strip() for h2 in soup.find_all('h2')][:5]
        
        images = soup.find_all('img')
        analysis['total_images'] = len(images)
        analysis['images_without_alt'] = len([img for img in images if not img.get('alt')])
        
        links = soup.find_all('a', href=True)
        for link in links:
            href = link['href']
            if href.startswith('http') and url not in href:
                analysis['external_links'] += 1
            elif href.startswith('/') or url in href:
                analysis['internal_links'] += 1
        
        text = soup.get_text()
        analysis['word_count'] = len(text.split())
        
        return analysis
        
    except Exception as e:
        return {'error': str(e)}

# ===========================
# 🤖 AI AUDIT GENERATION
# ===========================
def generate_audit_content(url, site_data, ahrefs_data, audit_type):
    """Generates structured audit data using Gemini"""
    
    try:
        model = genai.GenerativeModel("gemini-2.0-flash-exp")
        
        site_name = url.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
        
        if audit_type == "Basic":
            prompt = f"""
You are Claudio, an expert SEO auditor. Generate a BASIC SEO audit for: {site_name}

**SITE DATA:**
- Title: {site_data.get('title')}
- Meta Description: {site_data.get('meta_description', 'Missing')}
- H1 Tags: {len(site_data.get('h1_tags', []))} found
- Total Images: {site_data.get('total_images', 0)}
- Images without ALT: {site_data.get('images_without_alt', 0)}
- Internal Links: {site_data.get('internal_links', 0)}
- External Links: {site_data.get('external_links', 0)}

Generate a concise audit (max 600 words) with:

**Executive Summary** (2-3 paragraphs)

**Technical Findings:**
- Meta Tags Analysis
- Content Structure
- Image Optimization
- Internal Linking

**Key Recommendations** (prioritized list of 5-8 specific actions)

RULES:
- English only
- Only mention ACTUAL issues found
- Be specific with numbers
- Prioritize recommendations (Critical/High/Medium)
"""
        else:  # Full audit
            ahrefs_summary = "No Ahrefs data"
            if ahrefs_data:
                ahrefs_summary = f"""
- Domain Rating: {ahrefs_data['domain_rating']}
- Backlinks: {ahrefs_data['backlinks']:,}
- Referring Domains: {ahrefs_data['referring_domains']:,}
- Organic Keywords: {ahrefs_data['organic_keywords']:,}
- Organic Traffic: {ahrefs_data['organic_traffic']:,}/month
"""
            
            prompt = f"""
You are Claudio, expert SEO auditor. Generate a COMPREHENSIVE audit for: {site_name}

**SITE DATA:**
{site_data}

**AHREFS METRICS:**
{ahrefs_summary}

Generate a complete professional audit following this EXACT structure:

## EXECUTIVE SUMMARY
- Overall score /100
- 3-4 paragraph comprehensive summary
- Key metrics overview
- Main strengths and critical issues

## TECHNICAL SEO ANALYSIS
### Site Structure & Indexation
### Meta Tags & On-Page Elements
### Performance & Mobile

## BACKLINK PROFILE ANALYSIS
- Quality assessment
- Link diversity
- Toxic links (if any)
- Opportunities

## ORGANIC PERFORMANCE
### Current Rankings
### Top Performing Pages
### Keyword Opportunities

## QUICK WINS
(3-5 easy high-impact actions, 1-2 days each)

## COMPETITIVE ANALYSIS
(Brief competitive insights)

## PRIORITIZED ACTION PLAN

### CRITICAL (Week 1-2)
[Specific issues with:
- Description
- Impact (High/Medium/Low)
- Effort (hours)
- Expected result]

### HIGH PRIORITY (Week 3-4)
[Same format]

### MEDIUM PRIORITY (Month 2)
[Same format]

## EXPECTED RESULTS (3 Months)
[Realistic projections]

CRITICAL RULES:
- English only
- Only include ACTUAL issues found in the data
- NO placeholder issues
- Be specific with all numbers
- Use real Ahrefs data, don't invent
- If section has no issues, say "No critical issues found" and move on
"""
        
        response = model.generate_content(prompt)
        return response.text
        
    except Exception as e:
        return f"Error generating audit: {str(e)}"

# ===========================
# 📄 DOCUMENT GENERATION FROM TEMPLATE
# ===========================
def create_word_from_content(audit_content, site_name, audit_type):
    """Creates Word document from audit content"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'SEO Audit - {site_name}', 0)
    title.alignment = 1  # Center
    
    # Subtitle
    subtitle = doc.add_paragraph(f'{audit_type} Audit')
    subtitle.alignment = 1
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(96, 165, 250)
    
    doc.add_paragraph()  # Space
    
    # Process content
    lines = audit_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Headers
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        # Lists
        elif line.startswith(('- ', '* ')):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
        # Dividers
        elif line == '---':
            doc.add_paragraph('_' * 60)
        # Regular text
        else:
            line = line.replace('**', '')
            doc.add_paragraph(line)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Generated by Claudio - {datetime.now().strftime("%B %d, %Y")}')
    footer.alignment = 1
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def create_excel_from_content(audit_content, site_name):
    """Creates Excel task list from audit content"""
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "SEO Tasks"
    
    # Styling
    header_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    critical_fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
    high_fill = PatternFill(start_color="FED7AA", end_color="FED7AA", fill_type="solid")
    medium_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    low_fill = PatternFill(start_color="D1FAE5", end_color="D1FAE5", fill_type="solid")
    
    # Headers
    headers = ['#', 'Task', 'Category', 'Priority', 'Effort (hrs)', 'Impact', 'How to Fix', 'Status']
    ws.append(headers)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Extract tasks
    task_num = 1
    lines = audit_content.split('\n')
    current_priority = 'Medium'
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Detect priority
        if 'CRITICAL' in line.upper() and ('Week' in line or 'Priority' in line):
            current_priority = 'Critical'
        elif 'HIGH PRIORITY' in line.upper():
            current_priority = 'High'
        elif 'MEDIUM PRIORITY' in line.upper():
            current_priority = 'Medium'
        elif 'LOW PRIORITY' in line.upper():
            current_priority = 'Low'
        
        # Extract tasks
        if (line.startswith(('-', '*', '•')) or re.match(r'^\d+\.', line)) and len(line) > 30:
            # Check if in action section
            context = ' '.join(lines[max(0, i-15):i])
            if any(keyword in context for keyword in ['Action', 'Priority', 'Quick Win', 'Recommendation']):
                
                task = re.sub(r'^[\d\-\*\•\.\s]+', '', line).strip()
                
                # Estimate effort and impact
                effort = '2-4'
                impact = 'Medium'
                
                if current_priority == 'Critical':
                    effort = '4-8'
                    impact = 'High'
                elif current_priority == 'High':
                    effort = '3-6'
                    impact = 'High'
                elif current_priority == 'Low':
                    effort = '1-2'
                    impact = 'Low'
                
                # Categorize
                category = 'Technical SEO'
                task_lower = task.lower()
                if any(w in task_lower for w in ['content', 'keyword', 'text', 'copy', 'article']):
                    category = 'Content'
                elif any(w in task_lower for w in ['link', 'backlink', 'anchor']):
                    category = 'Link Building'
                elif any(w in task_lower for w in ['meta', 'title', 'description', 'tag', 'heading']):
                    category = 'On-Page SEO'
                elif any(w in task_lower for w in ['speed', 'performance', 'mobile', 'core web']):
                    category = 'Performance'
                
                # How to fix (extract from next few lines if available)
                how_to_fix = 'See audit report for details'
                
                row = [task_num, task, category, current_priority, effort, impact, how_to_fix, 'To Do']
                ws.append(row)
                
                # Color code by priority
                row_idx = ws.max_row
                fill = None
                if current_priority == 'Critical':
                    fill = critical_fill
                elif current_priority == 'High':
                    fill = high_fill
                elif current_priority == 'Medium':
                    fill = medium_fill
                elif current_priority == 'Low':
                    fill = low_fill
                
                if fill:
                    for cell in ws[row_idx]:
                        cell.fill = fill
                
                task_num += 1
    
    # Column widths
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 10
    ws.column_dimensions['G'].width = 35
    ws.column_dimensions['H'].width = 12
    
    # Save
    excel_io = BytesIO()
    wb.save(excel_io)
    excel_io.seek(0)
    
    return excel_io

# ===========================
# 🎨 SIDEBAR
# ===========================
with st.sidebar:
    st.markdown("### 🏢 System Status")
    
    if GEMINI_AVAILABLE:
        st.markdown('<span class="status-badge status-connected">🟢 Gemini Connected</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-disconnected">🔴 Gemini Offline</span>', unsafe_allow_html=True)
    
    if CLAUDE_AVAILABLE:
        st.markdown('<span class="status-badge status-connected">🟢 Claude Connected</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-disconnected">🔴 Claude Offline</span>', unsafe_allow_html=True)
    
    if AHREFS_AVAILABLE:
        st.markdown('<span class="status-badge status-connected">🟢 Ahrefs Connected</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="status-badge status-optional">⚠️ Ahrefs Optional</span>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    st.markdown("### ℹ️ About")
    st.markdown("""
    **Claudio** generates professional SEO audits in seconds.
    
    **Features**:
    - 🔍 Basic visual analysis
    - 💎 Full analysis with Ahrefs
    - 🤖 AI-powered insights
    - 📄 Professional reports
    """)
    
    st.markdown("---")
    st.caption("v2.1 - Template Edition")

# ===========================
# 🎯 MAIN INTERFACE
# ===========================

st.markdown("""
<div class="claudio-header">
    <div class="claudio-avatar-large">👔</div>
    <div class="claudio-title">CLAUDIO</div>
    <div class="claudio-subtitle">Professional SEO Auditor</div>
</div>
""", unsafe_allow_html=True)

# Configuration
col1, col2 = st.columns([2, 1])

with col1:
    audit_type = st.radio(
        "Audit Type",
        ["🔍 Basic (Visual Analysis)", "💎 Full (With Ahrefs Data)"],
        help="Basic: Quick visual analysis\nFull: Complete with Ahrefs metrics"
    )

with col2:
    if "Full" in audit_type:
        st.info("**Full Audit**\n\n✓ Domain Rating\n✓ Backlinks\n✓ Keywords\n✓ Traffic data")
    else:
        st.info("**Basic Audit**\n\n✓ Technical SEO\n✓ On-page analysis\n✓ Quick insights")

st.markdown("---")

# AI Model
col1, col2 = st.columns([3, 1])

with col1:
    available_models = []
    
    if GEMINI_AVAILABLE:
        available_models.append("⚡ Gemini 2.0 Flash")
    
    if CLAUDE_AVAILABLE:
        available_models.extend([
            "🎯 Claude Sonnet 4.5",
            "👑 Claude Opus 4.5"
        ])
    
    if not available_models:
        st.error("❌ No AI models configured.")
        st.stop()
    
    selected_model = st.selectbox("AI Model", available_models)

st.markdown("---")

# URL Input
url_input = st.text_input(
    "Website URL",
    placeholder="https://example.com",
    help="Enter the full URL including https://"
)

# Confirmation
if "Full" in audit_type:
    if AHREFS_AVAILABLE:
        st.warning("⚠️ Full Audit will use Ahrefs API credits")
        confirm_ahrefs = st.checkbox("✓ Confirm Ahrefs API usage", value=False)
    else:
        st.error("❌ Ahrefs API not configured. Cannot perform Full audit.")
        confirm_ahrefs = False
else:
    confirm_ahrefs = True

st.markdown("---")

# Generate Button
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    button_disabled = not url_input or not confirm_ahrefs
    
    if st.button("🚀 Generate Audit", disabled=button_disabled, use_container_width=True):
        
        if not url_input:
            st.error("❌ Please enter a URL")
        else:
            st.markdown("---")
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Step 1: Analyze site
            status_text.text("🔍 Analyzing website...")
            progress_bar.progress(20)
            site_data = analyze_basic_site(url_input)
            time.sleep(1)
            
            if 'error' in site_data:
                st.error(f"❌ Error: {site_data['error']}")
                st.stop()
            
            # Step 2: Get Ahrefs data if Full
            ahrefs_data = None
            if "Full" in audit_type and AHREFS_AVAILABLE:
                status_text.text("📊 Fetching Ahrefs data...")
                progress_bar.progress(40)
                ahrefs_data = get_ahrefs_data(url_input)
                time.sleep(1)
            
            # Step 3: Generate audit
            status_text.text("🤖 Generating audit content...")
            progress_bar.progress(60)
            
            type_audit = "Basic" if "Basic" in audit_type else "Full"
            
            audit_content = generate_audit_content(url_input, site_data, ahrefs_data, type_audit)
            
            # Step 4: Create documents
            status_text.text("📄 Creating documents...")
            progress_bar.progress(80)
            
            site_name = url_input.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
            
            doc_file = create_word_from_content(audit_content, site_name, type_audit)
            
            excel_file = None
            if type_audit == "Full":
                excel_file = create_excel_from_content(audit_content, site_name)
            
            progress_bar.progress(100)
            status_text.text("✅ Complete!")
            time.sleep(0.5)
            
            progress_bar.empty()
            status_text.empty()
            
            # Results
            st.markdown("---")
            st.success("✅ Audit completed successfully!")
            
            tab1, tab2 = st.tabs(["📄 Preview", "📥 Download"])
            
            with tab1:
                st.markdown('<div class="audit-report">', unsafe_allow_html=True)
                st.markdown(audit_content)
                st.markdown('</div>', unsafe_allow_html=True)
            
            with tab2:
                st.markdown("### Download Your Documents")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### 📄 Audit Report")
                    st.download_button(
                        label="📥 Download Report (.docx)",
                        data=doc_file,
                        file_name=f"SEO_Audit_{site_name}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                with col2:
                    if excel_file:
                        st.markdown("#### 📊 Task List")
                        st.download_button(
                            label="📥 Download Tasks (.xlsx)",
                            data=excel_file,
                            file_name=f"SEO_Tasks_{site_name}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    else:
                        st.info("📊 Task list only for Full audits")

# Footer
st.markdown("---")
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("**Claudio SEO Auditor**")
    st.caption("Professional audits in seconds")

with col2:
    st.markdown("**Powered by**")
    st.caption("Anthropic • Google • Ahrefs")

with col3:
    st.markdown("**Need help?**")
    st.caption("[Documentation](#) • [Support](#)")
